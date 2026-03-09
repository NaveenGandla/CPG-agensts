"""
Template parsing service — extracts structure from uploaded CPG template files.

Uses Azure Document Intelligence to extract text, headings, tables, and
structural elements from PDF/DOCX/PPTX files, then produces a template
specification the agent can use to generate CPGs in the same format.
"""

import io
import json
from typing import Any

from src.core.logging import get_logger
from src.core.settings import get_settings

logger = get_logger("template_parser")

SUPPORTED_EXTENSIONS = {"pdf", "docx", "pptx", "xlsx", "txt"}

CONTENT_TYPE_MAP = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "txt": "text/plain",
}


async def parse_template(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """
    Parse an uploaded template file and extract its structure.

    Returns a dict with:
      - filename: original filename
      - sections: list of detected section headings with descriptions
      - tables: list of table structures found
      - full_text: raw extracted text (for the agent to analyse)
      - template_instructions: formatted instruction string for the agent
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '.{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    settings = get_settings().azure_doc_intelligence

    if settings.endpoint and ext != "txt":
        result = await _parse_with_document_intelligence(file_bytes, ext, settings)
    else:
        result = _parse_with_fallback(file_bytes, ext)

    result["filename"] = filename
    result["template_instructions"] = _build_template_instructions(result)
    return result


async def _parse_with_document_intelligence(
    file_bytes: bytes, ext: str, settings: Any
) -> dict[str, Any]:
    """Use Azure Document Intelligence to extract structure."""
    from azure.ai.documentintelligence.aio import DocumentIntelligenceClient
    from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
    from azure.core.credentials import AzureKeyCredential
    from azure.identity.aio import DefaultAzureCredential

    if settings.api_key:
        credential = AzureKeyCredential(settings.api_key)
    else:
        credential = DefaultAzureCredential()

    content_type = CONTENT_TYPE_MAP.get(ext, "application/octet-stream")

    async with DocumentIntelligenceClient(
        endpoint=settings.endpoint, credential=credential
    ) as client:
        poller = await client.begin_analyze_document(
            model_id=settings.model_id,
            analyze_request=AnalyzeDocumentRequest(
                bytes_source=file_bytes,
            ),
            content_type=content_type,
        )
        di_result = await poller.result()

    sections: list[dict[str, str]] = []
    full_text_parts: list[str] = []
    tables: list[dict[str, Any]] = []

    # Extract paragraphs and detect headings
    if di_result.paragraphs:
        for para in di_result.paragraphs:
            role = getattr(para, "role", None)
            text = para.content.strip()
            if not text:
                continue
            full_text_parts.append(text)
            if role in ("title", "sectionHeading"):
                sections.append({"heading": text, "role": role or "sectionHeading"})

    # If no paragraphs, fall back to pages
    if not full_text_parts and di_result.pages:
        for page in di_result.pages:
            for line in (page.lines or []):
                full_text_parts.append(line.content)

    # Extract table structures
    if di_result.tables:
        for table in di_result.tables:
            headers = []
            for cell in (table.cells or []):
                if cell.row_index == 0:
                    headers.append(cell.content)
            tables.append({
                "rows": table.row_count,
                "columns": table.column_count,
                "headers": headers,
            })

    # If no headings detected from roles, infer from text patterns
    if not sections:
        sections = _infer_sections_from_text("\n".join(full_text_parts))

    logger.info(
        "DI parsed template: %d sections, %d tables, %d text chars",
        len(sections), len(tables), len("\n".join(full_text_parts)),
    )

    return {
        "sections": sections,
        "tables": tables,
        "full_text": "\n\n".join(full_text_parts),
    }


def _parse_with_fallback(file_bytes: bytes, ext: str) -> dict[str, Any]:
    """Fallback: extract text without Document Intelligence."""
    if ext == "txt":
        text = file_bytes.decode("utf-8", errors="replace")
    elif ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text = "\n\n".join(
            page.extract_text() or "" for page in reader.pages
        )
    elif ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        text = file_bytes.decode("utf-8", errors="replace")

    sections = _infer_sections_from_text(text)

    return {
        "sections": sections,
        "tables": [],
        "full_text": text,
    }


def _infer_sections_from_text(text: str) -> list[dict[str, str]]:
    """Heuristic: detect section headings from text patterns."""
    import re

    sections = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Numbered sections: "1. Title", "1.1 Title", "Section 1:"
        if re.match(r"^(\d+\.?\d*\.?\s+|Section\s+\d+)", line, re.IGNORECASE):
            sections.append({"heading": line, "role": "sectionHeading"})
        # ALL CAPS headings (at least 3 words, all uppercase)
        elif line.isupper() and len(line.split()) >= 2 and len(line) <= 100:
            sections.append({"heading": line, "role": "sectionHeading"})
        # Short lines ending with colon
        elif line.endswith(":") and len(line) <= 80 and len(line.split()) <= 8:
            sections.append({"heading": line.rstrip(":"), "role": "sectionHeading"})

    return sections


def _build_template_instructions(parsed: dict[str, Any]) -> str:
    """Convert parsed template structure into instructions for the agent."""
    parts = [
        "The user has uploaded a CPG template file. You MUST generate the CPG "
        "report following the EXACT structure and format of this template.",
        "",
        f"Template file: {parsed['filename']}",
    ]

    if parsed["sections"]:
        parts.append("")
        parts.append("## Detected Template Sections (in order):")
        for i, sec in enumerate(parsed["sections"], 1):
            parts.append(f"  {i}. {sec['heading']}")

    if parsed["tables"]:
        parts.append("")
        parts.append("## Table Structures Found:")
        for i, tbl in enumerate(parsed["tables"], 1):
            headers = ", ".join(tbl["headers"]) if tbl["headers"] else "no headers"
            parts.append(
                f"  Table {i}: {tbl['rows']} rows x {tbl['columns']} columns "
                f"(Headers: {headers})"
            )

    parts.append("")
    parts.append("## Full Template Content (for reference):")
    # Truncate if very long to stay within context limits
    full_text = parsed["full_text"]
    if len(full_text) > 8000:
        full_text = full_text[:8000] + "\n\n[... truncated for brevity ...]"
    parts.append(full_text)

    parts.append("")
    parts.append(
        "IMPORTANT: Reproduce the same section ordering, headings, and table "
        "formats from this template. Fill each section with clinically accurate "
        "content based on the evidence and topic provided."
    )

    return "\n".join(parts)
